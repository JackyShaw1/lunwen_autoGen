"""生成具有正式教学版式的 Word / PDF 授课包；PPTX 由独立课件服务负责。"""

from __future__ import annotations

import re
from datetime import datetime
from pathlib import Path

from docx import Document
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor

from app.config import settings
from app.services.material_service import get_cached_material_image, resolve_package_materials


_CN_FONT_CANDIDATES = [
    Path("/usr/share/fonts/truetype/wqy/wqy-microhei.ttc"),
    Path("/usr/share/fonts/truetype/wqy/wqy-zenhei.ttc"),
    Path("/usr/share/fonts/opentype/noto/NotoSansCJK-Regular.ttc"),
    Path("/usr/share/fonts/truetype/noto/NotoSansCJK-Regular.ttc"),
    Path(r"C:\Windows\Fonts\msyh.ttc"),
    Path(r"C:\Windows\Fonts\simsun.ttc"),
    Path(r"C:\Windows\Fonts\simhei.ttf"),
]

_INDIGO = "4F46E5"
_INDIGO_DARK = "312E81"
_INDIGO_LIGHT = "EEF2FF"
_SLATE = "475569"
_BORDER = "CBD5E1"


def _ensure_export_dir() -> Path:
    path = Path(settings.export_dir)
    path.mkdir(parents=True, exist_ok=True)
    return path


def safe_export_filename(title: str, version: int, file_format: str) -> str:
    """创建适合浏览器和不同操作系统的可读文件名。"""
    clean = re.sub(r"[\\/:*?\"<>|\x00-\x1f]", "_", (title or "教学案例").strip())
    clean = re.sub(r"\s+", " ", clean).strip(" ._")[:80] or "教学案例"
    return f"{clean}_教学案例授课包_V{version}.{file_format}"


def _split_reading_paragraphs(text: str, target_chars: int = 320) -> list[str]:
    """把模型可能返回的长文本切成适合屏幕和纸张阅读的自然段。"""
    normalized = str(text or "").replace("\r\n", "\n").strip()
    if not normalized:
        return ["（无）"]
    blocks = [part.strip() for part in re.split(r"\n\s*\n+", normalized) if part.strip()]
    result: list[str] = []
    for block in blocks:
        if len(block) <= target_chars * 1.35:
            result.append(block.replace("\n", " "))
            continue
        sentences = [s.strip() for s in re.split(r"(?<=[。！？；])", block) if s.strip()]
        current = ""
        for sentence in sentences:
            if current and len(current) + len(sentence) > target_chars:
                result.append(current)
                current = sentence
            else:
                current += sentence
        if current:
            result.append(current)
    return result or [normalized]


def _set_run_font(run, name: str = "微软雅黑", size: float | None = None, bold: bool | None = None) -> None:
    run.font.name = name
    run._element.get_or_add_rPr().rFonts.set(qn("w:eastAsia"), name)
    if size is not None:
        run.font.size = Pt(size)
    if bold is not None:
        run.bold = bold


def _shade_cell(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def _set_cell_border(cell, color: str = _BORDER) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    borders = tc_pr.first_child_found_in("w:tcBorders")
    if borders is None:
        borders = OxmlElement("w:tcBorders")
        tc_pr.append(borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        tag = f"w:{edge}"
        element = borders.find(qn(tag))
        if element is None:
            element = OxmlElement(tag)
            borders.append(element)
        element.set(qn("w:val"), "single")
        element.set(qn("w:sz"), "4")
        element.set(qn("w:color"), color)


def _add_page_number(paragraph) -> None:
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = paragraph.add_run("第 ")
    _set_run_font(run, size=9)
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    instruction = OxmlElement("w:instrText")
    instruction.set(qn("xml:space"), "preserve")
    instruction.text = " PAGE "
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    run._r.extend([begin, instruction, end])
    end_run = paragraph.add_run(" 页")
    _set_run_font(end_run, size=9)


def _add_docx_text(doc: Document, text: str, *, indent: bool = True) -> None:
    for block in _split_reading_paragraphs(text):
        paragraph = doc.add_paragraph()
        paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        paragraph.paragraph_format.line_spacing = 1.55
        paragraph.paragraph_format.space_after = Pt(8)
        if indent:
            paragraph.paragraph_format.first_line_indent = Pt(22)
        run = paragraph.add_run(block)
        _set_run_font(run, "宋体", 11)


def _add_docx_visual_materials(doc: Document, package: dict) -> None:
    assets = resolve_package_materials(package)
    if not assets:
        return
    doc.add_heading("官方视觉材料（教学引用）", level=2)
    intro = doc.add_paragraph("以下图片来自标注的政府或机构官网，用于课堂教学引用；公开传播或商业再利用前请确认授权。")
    for run in intro.runs:
        _set_run_font(run, "微软雅黑", 9)
        run.font.color.rgb = RGBColor.from_string("64748B")
    for index, asset in enumerate(assets, 1):
        try:
            image_path = get_cached_material_image(asset["id"])
            image_paragraph = doc.add_paragraph()
            image_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            image_paragraph.add_run().add_picture(str(image_path), width=Cm(14.8))
        except Exception:  # Export text and provenance even if the official host is temporarily unavailable.
            unavailable = doc.add_paragraph("图片暂时无法获取，请通过原始来源页面查看。")
            unavailable.alignment = WD_ALIGN_PARAGRAPH.CENTER
        caption = doc.add_paragraph()
        caption.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = caption.add_run(f"图 {index}　{asset['title']}｜{asset['caption']}")
        _set_run_font(run, "微软雅黑", 9, True)
        source = doc.add_paragraph()
        source.alignment = WD_ALIGN_PARAGRAPH.CENTER
        source_run = source.add_run(
            f"来源：{asset['source_org']}　摄影：{asset.get('photographer') or '未标注'}　{asset['source_page_url']}"
        )
        _set_run_font(source_run, "微软雅黑", 8)
        source_run.font.color.rgb = RGBColor.from_string("64748B")


def _add_docx_evidence_sources(doc: Document, package: dict) -> None:
    sources = package.get("evidence_sources") or []
    if not sources:
        return
    doc.add_heading("八、事实来源与引用说明", level=1)
    intro = doc.add_paragraph("正文中的 [S1]、[S2] 等标记与下列来源一一对应。课堂讲授时应区分公开事实、分析推断与教学任务，不补写未公开数据或人物对话。")
    for run in intro.runs:
        _set_run_font(run, "微软雅黑", 9)
        run.font.color.rgb = RGBColor.from_string("64748B")
    for source in sources:
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(8)
        label = p.add_run(f"[{source.get('id', '')}] {source.get('title', '')}\n")
        _set_run_font(label, "微软雅黑", 10, True)
        detail = p.add_run(
            f"{source.get('source_org', '')}"
            f"{f' · {source.get('published_at')}' if source.get('published_at') else ''}\n"
            f"用途：{source.get('usage', '')}\n{source.get('source_page_url', '')}"
        )
        _set_run_font(detail, "宋体", 9)
        detail.font.color.rgb = RGBColor.from_string("475569")


def _configure_docx(doc: Document, course: str) -> None:
    section = doc.sections[0]
    section.top_margin = Cm(2.2)
    section.bottom_margin = Cm(2)
    section.left_margin = Cm(2.5)
    section.right_margin = Cm(2.3)

    normal = doc.styles["Normal"]
    normal.font.name = "宋体"
    normal.font.size = Pt(11)
    normal._element.get_or_add_rPr().rFonts.set(qn("w:eastAsia"), "宋体")
    normal.paragraph_format.line_spacing = 1.5
    normal.paragraph_format.space_after = Pt(6)

    for style_name, size, color in (("Title", 24, _INDIGO_DARK), ("Heading 1", 16, _INDIGO_DARK), ("Heading 2", 13, _SLATE)):
        style = doc.styles[style_name]
        style.font.name = "微软雅黑"
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = RGBColor.from_string(color)
        style._element.get_or_add_rPr().rFonts.set(qn("w:eastAsia"), "微软雅黑")
        style.paragraph_format.space_before = Pt(14)
        style.paragraph_format.space_after = Pt(8)

    header = section.header.paragraphs[0]
    header.text = f"知案 · AI 教学案例工作台    |    {course or '教学案例授课包'}"
    header.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    for run in header.runs:
        _set_run_font(run, "微软雅黑", 8)
        run.font.color.rgb = RGBColor.from_string("94A3B8")
    _add_page_number(section.footer.paragraphs[0])


def export_docx(package: dict, title: str, version: int = 1) -> str:
    meta = package.get("meta", {}) or {}
    body = package.get("body", {}) or {}
    guide = package.get("instructor_guide", {}) or {}
    quality = package.get("quality", {}) or {}
    display_title = meta.get("title") or title
    filename = safe_export_filename(display_title, version, "docx")
    path = _ensure_export_dir() / filename

    doc = Document()
    _configure_docx(doc, meta.get("course", ""))
    doc.core_properties.title = display_title
    doc.core_properties.subject = "教学案例授课包"
    doc.core_properties.author = "知案 · AI 教学案例工作台"

    # 封面
    doc.add_paragraph("\n\n")
    kicker = doc.add_paragraph("教学案例授课包")
    kicker.alignment = WD_ALIGN_PARAGRAPH.CENTER
    kicker_run = kicker.runs[0]
    _set_run_font(kicker_run, "微软雅黑", 12, True)
    kicker_run.font.color.rgb = RGBColor.from_string(_INDIGO)
    title_p = doc.add_paragraph(style="Title")
    title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_p.paragraph_format.space_before = Pt(28)
    title_run = title_p.add_run(display_title)
    _set_run_font(title_run, "微软雅黑", 24, True)

    meta_table = doc.add_table(rows=3, cols=2)
    meta_table.alignment = 1
    cover_rows = [
        ("课程名称", meta.get("course", "—"), "学科领域", meta.get("subject", "—")),
        ("案例类型", meta.get("case_type", "—"), "适用难度", meta.get("difficulty", "—")),
        ("文档版本", f"V{version}", "生成日期", datetime.now().strftime("%Y-%m-%d")),
    ]
    for row, values in zip(meta_table.rows, cover_rows):
        left_label, left_value, right_label, right_value = values
        row.cells[0].text = f"{left_label}\n{left_value}"
        row.cells[1].text = f"{right_label}\n{right_value}"
        for cell in row.cells:
            _shade_cell(cell, "F8FAFC")
            _set_cell_border(cell, "E2E8F0")
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            for paragraph in cell.paragraphs:
                paragraph.paragraph_format.space_after = Pt(2)
                for run in paragraph.runs:
                    _set_run_font(run, "微软雅黑", 10)

    disclaimer = doc.add_paragraph()
    disclaimer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    disclaimer.paragraph_format.space_before = Pt(20)
    run = disclaimer.add_run(meta.get("fictional_disclaimer") or "本材料仅用于教学研讨。")
    _set_run_font(run, "微软雅黑", 9)
    run.font.color.rgb = RGBColor.from_string("64748B")
    doc.add_page_break()

    doc.add_heading("一、学习目标", level=1)
    for index, objective in enumerate(package.get("learning_objectives") or [], 1):
        p = doc.add_paragraph()
        p.paragraph_format.left_indent = Cm(0.35)
        p.paragraph_format.first_line_indent = Cm(-0.35)
        label = p.add_run(f"{index:02d}  ")
        _set_run_font(label, "微软雅黑", 10, True)
        label.font.color.rgb = RGBColor.from_string(_INDIGO)
        content = p.add_run(objective.get("description", ""))
        _set_run_font(content, "宋体", 11)

    doc.add_heading("二、案例背景", level=1)
    _add_docx_text(doc, body.get("background", ""))
    _add_docx_visual_materials(doc, package)

    doc.add_heading("三、案例正文", level=1)
    _add_docx_text(doc, body.get("narrative", ""))

    doc.add_heading("四、关键决策点", level=1)
    callout = doc.add_table(rows=1, cols=1)
    cell = callout.cell(0, 0)
    _shade_cell(cell, _INDIGO_LIGHT)
    _set_cell_border(cell, "C7D2FE")
    cell.text = ""
    paragraph = cell.paragraphs[0]
    paragraph.paragraph_format.space_before = Pt(8)
    paragraph.paragraph_format.space_after = Pt(8)
    run = paragraph.add_run(body.get("decision_point", "") or "（无）")
    _set_run_font(run, "微软雅黑", 11, True)
    run.font.color.rgb = RGBColor.from_string(_INDIGO_DARK)

    characters = body.get("characters") or []
    if characters:
        doc.add_heading("五、角色与立场", level=1)
        table = doc.add_table(rows=1, cols=3)
        table.style = "Table Grid"
        headers = ("角色", "身份", "核心立场")
        for cell, header_text in zip(table.rows[0].cells, headers):
            cell.text = header_text
            _shade_cell(cell, _INDIGO_DARK)
            for run in cell.paragraphs[0].runs:
                _set_run_font(run, "微软雅黑", 10, True)
                run.font.color.rgb = RGBColor(255, 255, 255)
        for character in characters:
            cells = table.add_row().cells
            for cell, value in zip(cells, (character.get("name", ""), character.get("role", ""), character.get("stance", ""))):
                cell.text = str(value)
                for run in cell.paragraphs[0].runs:
                    _set_run_font(run, "宋体", 10)

    doc.add_heading("六、课堂讨论题", level=1)
    for index, question in enumerate(package.get("discussion_questions") or [], 1):
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(7)
        number = p.add_run(f"Q{index}  ")
        _set_run_font(number, "微软雅黑", 10, True)
        number.font.color.rgb = RGBColor.from_string(_INDIGO)
        level = p.add_run(f"[{question.get('level', '')}] ")
        _set_run_font(level, "微软雅黑", 9, True)
        level.font.color.rgb = RGBColor.from_string(_SLATE)
        text_run = p.add_run(question.get("question", ""))
        _set_run_font(text_run, "宋体", 11, True)
        if question.get("teaching_intent"):
            hint = doc.add_paragraph(f"教学意图：{question['teaching_intent']}")
            hint.paragraph_format.left_indent = Cm(0.7)
            for hint_run in hint.runs:
                _set_run_font(hint_run, "微软雅黑", 9)
                hint_run.font.color.rgb = RGBColor.from_string("64748B")

    doc.add_heading("七、教师授课指南", level=1)
    doc.add_heading("建议流程", level=2)
    _add_docx_text(doc, guide.get("teaching_flow", ""), indent=False)
    for label, items in (("教学要点", guide.get("key_points") or []), ("常见误区", guide.get("common_misconceptions") or [])):
        if items:
            doc.add_heading(label, level=2)
            for item in items:
                p = doc.add_paragraph(style="List Bullet")
                run = p.add_run(str(item))
                _set_run_font(run, "宋体", 11)

    _add_docx_evidence_sources(doc, package)

    matrix = package.get("alignment_matrix") or []
    if matrix:
        doc.add_heading("九、教学目标对齐表", level=1)
        table = doc.add_table(rows=1, cols=4)
        table.style = "Table Grid"
        for cell, header_text in zip(table.rows[0].cells, ("教学目标", "案例环节", "课堂活动", "评价方式")):
            cell.text = header_text
            _shade_cell(cell, _INDIGO_DARK)
            for run in cell.paragraphs[0].runs:
                _set_run_font(run, "微软雅黑", 9, True)
                run.font.color.rgb = RGBColor(255, 255, 255)
        for row in matrix:
            cells = table.add_row().cells
            values = (row.get("objective_id", ""), row.get("case_section", ""), row.get("activity", ""), row.get("assessment", ""))
            for cell, value in zip(cells, values):
                cell.text = str(value)
                for run in cell.paragraphs[0].runs:
                    _set_run_font(run, "宋体", 9)

    if quality:
        doc.add_heading("十、质量评审", level=1)
        score = doc.add_paragraph()
        score_run = score.add_run(f"综合评分  {quality.get('overall_score', '—')} / 5")
        _set_run_font(score_run, "微软雅黑", 13, True)
        score_run.font.color.rgb = RGBColor.from_string(_INDIGO)
        _add_docx_text(doc, quality.get("reviewer_summary", "") or "（无）", indent=False)

    doc.save(str(path))
    return str(path)


def _resolve_cn_font() -> Path | None:
    return next((path for path in _CN_FONT_CANDIDATES if path.exists()), None)


def _register_cn_font() -> str:
    from reportlab.pdfbase import pdfmetrics
    from reportlab.pdfbase.cidfonts import UnicodeCIDFont
    from reportlab.pdfbase.ttfonts import TTFont

    font_name = "CaseAutoGenCN"
    if font_name in pdfmetrics.getRegisteredFontNames():
        return font_name
    path = _resolve_cn_font()
    if path:
        kwargs = {"subfontIndex": 0} if path.suffix.lower() == ".ttc" else {}
        pdfmetrics.registerFont(TTFont(font_name, str(path), **kwargs))
        return font_name
    # 最后使用 ReportLab 自带的中文 CID 字体，避免环境差异导致整个 PDF 功能不可用。
    pdfmetrics.registerFont(UnicodeCIDFont("STSong-Light"))
    return "STSong-Light"


def _pdf_escape(text: str) -> str:
    return str(text or "").replace("&", "&amp;").replace("<", "&lt;").replace(">", "&gt;")


def export_pdf(package: dict, title: str, version: int = 1) -> str:
    from reportlab.lib import colors
    from reportlab.lib.enums import TA_CENTER, TA_JUSTIFY, TA_LEFT
    from reportlab.lib.pagesizes import A4
    from reportlab.lib.styles import ParagraphStyle, getSampleStyleSheet
    from reportlab.lib.units import cm
    from reportlab.platypus import Image, KeepTogether, PageBreak, Paragraph, SimpleDocTemplate, Spacer, Table, TableStyle

    meta = package.get("meta", {}) or {}
    body = package.get("body", {}) or {}
    guide = package.get("instructor_guide", {}) or {}
    quality = package.get("quality", {}) or {}
    display_title = meta.get("title") or title
    filename = safe_export_filename(display_title, version, "pdf")
    pdf_path = _ensure_export_dir() / filename
    font_name = _register_cn_font()
    styles = getSampleStyleSheet()

    title_style = ParagraphStyle("CNTitle", parent=styles["Title"], fontName=font_name, fontSize=24, leading=34, alignment=TA_CENTER, textColor=colors.HexColor("#312e81"), wordWrap="CJK")
    kicker_style = ParagraphStyle("CNKicker", parent=styles["Normal"], fontName=font_name, fontSize=11, leading=18, alignment=TA_CENTER, textColor=colors.HexColor("#4f46e5"), spaceAfter=18)
    h1_style = ParagraphStyle("CNH1", parent=styles["Heading1"], fontName=font_name, fontSize=15, leading=23, spaceBefore=16, spaceAfter=9, textColor=colors.HexColor("#312e81"), keepWithNext=True)
    h2_style = ParagraphStyle("CNH2", parent=styles["Heading2"], fontName=font_name, fontSize=11, leading=18, spaceBefore=9, spaceAfter=5, textColor=colors.HexColor("#475569"), keepWithNext=True)
    body_style = ParagraphStyle("CNBody", parent=styles["Normal"], fontName=font_name, fontSize=10.5, leading=19, alignment=TA_JUSTIFY, firstLineIndent=21, spaceAfter=9, wordWrap="CJK", textColor=colors.HexColor("#1e293b"))
    plain_style = ParagraphStyle("CNPlain", parent=body_style, firstLineIndent=0)
    tip_style = ParagraphStyle("CNTip", parent=plain_style, fontSize=8.5, leading=14, textColor=colors.HexColor("#64748b"))
    table_style = ParagraphStyle("CNTable", parent=plain_style, fontSize=8.5, leading=13)
    question_style = ParagraphStyle("CNQuestion", parent=plain_style, fontSize=10.5, leading=18, spaceBefore=7, spaceAfter=3, textColor=colors.HexColor("#1e293b"))

    def paragraph_blocks(text: str, style=body_style):
        return [Paragraph(_pdf_escape(block).replace("\n", "<br/>"), style) for block in _split_reading_paragraphs(text)]

    story: list = [
        Spacer(1, 2.7 * cm),
        Paragraph("教学案例授课包", kicker_style),
        Paragraph(_pdf_escape(display_title), title_style),
        Spacer(1, 1.2 * cm),
    ]
    cover_data = [
        ["课程名称", meta.get("course", "—"), "学科领域", meta.get("subject", "—")],
        ["案例类型", meta.get("case_type", "—"), "适用难度", meta.get("difficulty", "—")],
        ["文档版本", f"V{version}", "生成日期", datetime.now().strftime("%Y-%m-%d")],
    ]
    cover_table = Table(cover_data, colWidths=[2.3 * cm, 4.8 * cm, 2.3 * cm, 4.8 * cm], rowHeights=1.15 * cm)
    cover_table.setStyle(TableStyle([
        ("FONTNAME", (0, 0), (-1, -1), font_name), ("FONTSIZE", (0, 0), (-1, -1), 9),
        ("TEXTCOLOR", (0, 0), (-1, -1), colors.HexColor("#334155")),
        ("BACKGROUND", (0, 0), (-1, -1), colors.HexColor("#f8fafc")),
        ("BOX", (0, 0), (-1, -1), .5, colors.HexColor("#cbd5e1")),
        ("INNERGRID", (0, 0), (-1, -1), .35, colors.HexColor("#e2e8f0")),
        ("VALIGN", (0, 0), (-1, -1), "MIDDLE"), ("LEFTPADDING", (0, 0), (-1, -1), 7),
    ]))
    story.extend([cover_table, Spacer(1, .8 * cm)])
    if meta.get("fictional_disclaimer"):
        story.append(Paragraph(_pdf_escape(meta["fictional_disclaimer"]), tip_style))
    story.append(PageBreak())

    story.append(Paragraph("一、学习目标", h1_style))
    objectives = package.get("learning_objectives") or []
    for index, objective in enumerate(objectives, 1):
        story.append(Paragraph(_pdf_escape(f"{index:02d}　{objective.get('description', '')}"), plain_style))

    story.append(Paragraph("二、案例背景", h1_style))
    story.extend(paragraph_blocks(body.get("background", "")))
    visual_assets = resolve_package_materials(package)
    if visual_assets:
        story.append(Paragraph("官方视觉材料（教学引用）", h2_style))
        story.append(Paragraph("图片来自标注的政府或机构官网；公开传播或商业再利用前请确认授权。", tip_style))
        for index, asset in enumerate(visual_assets, 1):
            try:
                image_path = get_cached_material_image(asset["id"])
                image = Image(str(image_path), width=15.2 * cm, height=9.2 * cm, kind="proportional")
                caption = Paragraph(_pdf_escape(f"图 {index}　{asset['title']}｜{asset['caption']}"), table_style)
                source = Paragraph(
                    _pdf_escape(f"来源：{asset['source_org']}　摄影：{asset.get('photographer') or '未标注'}　{asset['source_page_url']}"),
                    tip_style,
                )
                story.append(KeepTogether([Spacer(1, .25 * cm), image, Spacer(1, .12 * cm), caption, source]))
            except Exception:
                story.append(Paragraph(_pdf_escape(f"{asset['title']}（图片暂时无法获取）｜来源：{asset['source_page_url']}"), tip_style))
    story.append(Paragraph("三、案例正文", h1_style))
    story.extend(paragraph_blocks(body.get("narrative", "")))
    story.append(Paragraph("四、关键决策点", h1_style))
    decision_table = Table([[Paragraph(_pdf_escape(body.get("decision_point", "") or "（无）"), plain_style)]], colWidths=[16.5 * cm])
    decision_table.setStyle(TableStyle([
        ("BACKGROUND", (0, 0), (-1, -1), colors.HexColor("#eef2ff")),
        ("BOX", (0, 0), (-1, -1), .7, colors.HexColor("#c7d2fe")),
        ("LEFTPADDING", (0, 0), (-1, -1), 12), ("RIGHTPADDING", (0, 0), (-1, -1), 12),
        ("TOPPADDING", (0, 0), (-1, -1), 10), ("BOTTOMPADDING", (0, 0), (-1, -1), 10),
    ]))
    story.append(decision_table)

    characters = body.get("characters") or []
    if characters:
        story.append(Paragraph("五、角色与立场", h1_style))
        character_data = [["角色", "身份", "核心立场"]] + [[c.get("name", ""), c.get("role", ""), c.get("stance", "")] for c in characters]
        table = Table(character_data, colWidths=[3 * cm, 4 * cm, 9.5 * cm], repeatRows=1)
        table.setStyle(TableStyle([
            ("FONTNAME", (0, 0), (-1, -1), font_name), ("FONTSIZE", (0, 0), (-1, -1), 8.5),
            ("BACKGROUND", (0, 0), (-1, 0), colors.HexColor("#312e81")),
            ("TEXTCOLOR", (0, 0), (-1, 0), colors.white),
            ("GRID", (0, 0), (-1, -1), .4, colors.HexColor("#cbd5e1")),
            ("VALIGN", (0, 0), (-1, -1), "TOP"), ("LEFTPADDING", (0, 0), (-1, -1), 6),
            ("TOPPADDING", (0, 0), (-1, -1), 6), ("BOTTOMPADDING", (0, 0), (-1, -1), 6),
        ]))
        story.append(table)

    story.append(Paragraph("六、课堂讨论题", h1_style))
    for index, question in enumerate(package.get("discussion_questions") or [], 1):
        story.append(Paragraph(_pdf_escape(f"Q{index}　[{question.get('level', '')}] {question.get('question', '')}"), question_style))
        if question.get("teaching_intent"):
            story.append(Paragraph(_pdf_escape(f"教学意图：{question['teaching_intent']}"), tip_style))

    story.append(Paragraph("七、教师授课指南", h1_style))
    story.append(Paragraph("建议流程", h2_style))
    story.extend(paragraph_blocks(guide.get("teaching_flow", ""), plain_style))
    for label, items in (("教学要点", guide.get("key_points") or []), ("常见误区", guide.get("common_misconceptions") or [])):
        if items:
            story.append(Paragraph(label, h2_style))
            for item in items:
                story.append(Paragraph(_pdf_escape(f"•　{item}"), plain_style))

    evidence_sources = package.get("evidence_sources") or []
    if evidence_sources:
        story.append(Paragraph("八、事实来源与引用说明", h1_style))
        story.append(Paragraph("正文中的 [S1]、[S2] 等标记与下列来源一一对应；不得补写未公开数据或人物对话。", tip_style))
        for source in evidence_sources:
            source_text = (
                f"[{source.get('id', '')}] {source.get('title', '')}<br/>"
                f"{source.get('source_org', '')}"
                f"{' · ' + str(source.get('published_at')) if source.get('published_at') else ''}<br/>"
                f"用途：{source.get('usage', '')}<br/>{source.get('source_page_url', '')}"
            )
            story.append(Paragraph(_pdf_escape(source_text).replace("&lt;br/&gt;", "<br/>"), table_style))

    matrix = package.get("alignment_matrix") or []
    if matrix:
        story.append(Paragraph("九、教学目标对齐表", h1_style))
        data = [[Paragraph(text, table_style) for text in ("教学目标", "案例环节", "课堂活动", "评价方式")]]
        for row in matrix:
            data.append([Paragraph(_pdf_escape(str(row.get(key, ""))), table_style) for key in ("objective_id", "case_section", "activity", "assessment")])
        table = Table(data, colWidths=[2.4 * cm, 5 * cm, 4 * cm, 5.1 * cm], repeatRows=1)
        table.setStyle(TableStyle([
            ("FONTNAME", (0, 0), (-1, -1), font_name),
            ("BACKGROUND", (0, 0), (-1, 0), colors.HexColor("#312e81")),
            ("TEXTCOLOR", (0, 0), (-1, 0), colors.white),
            ("GRID", (0, 0), (-1, -1), .4, colors.HexColor("#cbd5e1")),
            ("VALIGN", (0, 0), (-1, -1), "TOP"), ("LEFTPADDING", (0, 0), (-1, -1), 5),
            ("TOPPADDING", (0, 0), (-1, -1), 6), ("BOTTOMPADDING", (0, 0), (-1, -1), 6),
        ]))
        story.append(table)

    if quality:
        story.append(Paragraph("十、质量评审", h1_style))
        story.append(Paragraph(_pdf_escape(f"综合评分　{quality.get('overall_score', '—')} / 5"), h2_style))
        story.extend(paragraph_blocks(quality.get("reviewer_summary", "") or "（无）", plain_style))

    def draw_page(canvas, doc):
        canvas.saveState()
        canvas.setFont(font_name, 8)
        canvas.setFillColor(colors.HexColor("#94a3b8"))
        canvas.drawString(2 * cm, 1.05 * cm, f"知案 · {meta.get('course', '教学案例授课包')}")
        canvas.drawRightString(A4[0] - 2 * cm, 1.05 * cm, f"第 {doc.page} 页  ·  V{version}")
        canvas.setStrokeColor(colors.HexColor("#e2e8f0"))
        canvas.line(2 * cm, 1.35 * cm, A4[0] - 2 * cm, 1.35 * cm)
        canvas.restoreState()

    doc = SimpleDocTemplate(
        str(pdf_path), pagesize=A4, leftMargin=2.25 * cm, rightMargin=2.25 * cm,
        topMargin=1.8 * cm, bottomMargin=1.8 * cm, title=display_title,
        author="知案 · AI 教学案例工作台", subject="教学案例授课包",
    )
    doc.build(story, onFirstPage=draw_page, onLaterPages=draw_page)
    return str(pdf_path)
